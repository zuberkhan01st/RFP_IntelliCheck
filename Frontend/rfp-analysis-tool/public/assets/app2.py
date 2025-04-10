import os
import getpass
import time
import nest_asyncio
from langchain.embeddings import HuggingFaceHubEmbeddings
from dotenv import load_dotenv
from llama_parse import LlamaParse
from langchain.chains import RetrievalQA
from langchain.memory import ConversationBufferMemory
from langchain_groq import ChatGroq
import json
# Updated import for HuggingFaceHubEmbeddings
# from langchain_huggingface import HuggingFaceHubEmbeddings
from langgraph.graph import StateGraph, END
from langgraph.prebuilt import ToolInvocation, ToolExecutor
from typing import Dict, List, Tuple, Union
from langchain_core.messages import BaseMessage, FunctionMessage, HumanMessage
from langchain_core.runnables import RunnableConfig
from pydantic import BaseModel, Field
from langchain.docstore.document import Document
import docx
import chromadb
# Updated import for Chroma
from langchain_community.vectorstores import Chroma

nest_asyncio.apply()
load_dotenv()

groq_api_key = os.environ.get("GROQ_API_KEY")
llama_parse_api_key = os.environ.get("LLAMA_PARSE_API_KEY")
huggingfacehub_api_token = os.environ.get("HUGGINGFACEHUB_API_TOKEN")

if not all([groq_api_key, llama_parse_api_key, huggingfacehub_api_token]):
    raise ValueError("One or more API keys are missing in your .env file.")

parser = LlamaParse(
    result_type="markdown",
    api_key=llama_parse_api_key,
)

model_name = "sentence-transformers/all-mpnet-base-v2"
hf_embeddings = HuggingFaceHubEmbeddings(
    model=model_name,
    task="feature-extraction",
    huggingfacehub_api_token=huggingfacehub_api_token,
)

llm = ChatGroq(api_key = groq_api_key, model_name = 'llama-3.3-70b-versatile')

def setup_chroma_vector_store(embedding, collection_name="rfp_collection"):
    client = chromadb.Client()
    vector_store = Chroma(
        collection_name=collection_name,
        embedding_function=embedding,
        client=client
    )
    print(f"Chroma vector store initialized with collection: {collection_name}")
    return vector_store

def parse_document_llama_parse(file_path):
    try:
        if file_path.lower().endswith(".pdf"):
            print("Parsing PDF document with LlamaParse (metadata discarded)...")
            documents_from_llama_parse = parser.load_data(file_path=file_path)
            full_document_text = ""
            for doc in documents_from_llama_parse:
                full_document_text += doc.text + "\n\n"
            return full_document_text.strip()
        elif file_path.lower().endswith(".docx"):
            print("Parsing DOCX document with python-docx...")
            doc = docx.Document(file_path)
            text_list = [paragraph.text for paragraph in doc.paragraphs]
            return "\n\n".join(text_list)
        else:
            print(f"Unsupported file type: {file_path}. Only PDF and DOCX files are supported.")
            return None
    except Exception as e:
        print(f"Error parsing document: {e}")
        return None

def parse_docx_company_data(company_data_file):
    company_data = {}
    try:
        doc = docx.Document(company_data_file)
        for paragraph in doc.paragraphs:
            if ":" in paragraph.text:
                key, value = paragraph.text.split(":", 1)
                company_data[key.strip().lower().replace(" ", "_")] = value.strip()
        return company_data
    except Exception as e:
        print(f"Error parsing company data DOCX: {e}")
        return {}

def embed_and_store_in_chroma(vector_store, document_text):
    if not document_text:
        print("No document text to embed.")
        return

    chunk_size = 1000
    texts = [document_text[i:i + chunk_size] for i in range(0, len(document_text), chunk_size)]

    print(f"Embedding {len(texts)} document chunks and storing in Chroma")
    print("First chunk being sent to Chroma:", texts[0][:100] + "..." if len(texts[0]) > 100 else texts[0])

    try:
        vector_store.add_texts(
            texts=texts,
            metadatas=None
        )
        print("Document chunks successfully embedded and stored in Chroma.")
    except Exception as e:
        print(f"Error storing in Chroma: {e}")
        return None

    return vector_store

def load_company_data(company_data_file):
    if company_data_file.lower().endswith(".docx"):
        print("Loading company data from DOCX...")
        return parse_docx_company_data(company_data_file)
    else:
        print(f"Unsupported company data file type: {company_data_file}. Only DOCX is supported.")
        return None

def create_retrieval_qa_chain(vectorstore):
    return RetrievalQA.from_chain_type(
        llm=llm,
        retriever=vectorstore.as_retriever(search_kwargs={'k': 3}),
        chain_type="stuff"
    )

def check_eligibility(state):
    """Tool function to check eligibility."""
    company_data = state.company_data
    eligibility_qa_chain = state.eligibility_qa_chain
    
    prompt = f"""
    You are an Eligibility Assessment Specialist for RFPs (Request for Proposals). Your task is to determine if the company meets all eligibility requirements specified in the RFP document.
    
    Company Profile:
    {json.dumps(company_data, indent=2)}
    
    Based on the RFP document, please analyze:
    
    1. Does the company have the required years of experience in the relevant domain?
    2. Does the company possess all mandatory certifications mentioned in the RFP?
    3. Does the company have the necessary financial standing (revenue, turnover) as required?
    4. Does the company have experience with similar projects as specified in the RFP?
    5. Does the company meet geographical or jurisdictional requirements?
    
    For each criterion, provide:
    - ✅ Status: MET or NOT MET
    - 🔍 Evidence: Reference specific data points from company profile
    - 📋 Gap Analysis: If not met, what specifically is missing
    
    Conclude with:
    - Overall Eligibility: ELIGIBLE or NOT ELIGIBLE
    - Critical gaps requiring immediate attention
    - Recommendations for enhancing eligibility position
    """
    
    # Use the retrieval QA chain to get relevant RFP sections and analyze them
    response = eligibility_qa_chain.invoke({"query": prompt})
    return response["result"]

def generate_checklist(state):
    """Tool function to generate checklist items."""
    company_data = state.company_data
    checklist_qa_chain = state.checklist_qa_chain
    
    prompt = f"""
    You are an RFP Submission Specialist. Your task is to create a comprehensive checklist of all documents, forms, and requirements needed for a complete RFP submission.
    
    Company Profile:
    {json.dumps(company_data, indent=2)}
    
    Based on the RFP document, please create:
    
    1. 📋 DOCUMENT CHECKLIST:
       - List ALL required documents mentioned in the RFP
       - Include all forms, certificates, proofs, and attachments
       - Format as a numbered checklist with due dates where applicable
    
    2. 🔍 VERIFICATION STEPS:
       - For each document, provide verification criteria
       - Specify who should review/sign each document
       - Note any special requirements (notarization, apostille, etc.)
    
    3. ⚠️ COMMON PITFALLS:
       - Highlight frequently missed or incorrectly submitted items
       - Note any complex requirements that need special attention
    
    4. 📅 SUBMISSION TIMELINE:
       - Create a sequential timeline for gathering all materials
       - Allow sufficient time for approvals and signatures
       - Include buffer time before the final deadline
    
    Format the response as an actionable checklist that can be used by the submission team.
    """
    
    # Use the retrieval QA chain to get relevant RFP sections and analyze them
    response = checklist_qa_chain.invoke({"query": prompt})
    return response["result"]

def analyze_risk(state):
    """Tool function to analyze contract risks."""
    company_data = state.company_data
    risk_qa_chain = state.risk_qa_chain
    
    prompt = f"""
    You are a Legal Risk Analyst specializing in government and commercial RFPs. Your task is to identify and assess potential contractual and legal risks in the RFP.
    
    Company Profile and Risk Tolerance:
    {json.dumps(company_data, indent=2)}
    
    Based on the RFP document, please:
    
    1. 🚩 IDENTIFY HIGH-RISK CLAUSES:
       - Liability provisions (unlimited liability, broad indemnification)
       - Warranty and guarantee requirements
       - Intellectual Property rights transfer or licensing
       - Termination clauses (especially unilateral termination)
       - Payment terms and penalties
    
    2. ⚖️ RISK ASSESSMENT:
       - Categorize each identified clause as HIGH, MEDIUM, or LOW risk
       - Explain the potential impact on the company
       - Compare with industry standards for similar contracts
    
    3. 📝 MITIGATION STRATEGIES:
       - Suggest alternate clause language where appropriate
       - Recommend risk management approaches
       - Identify items requiring legal counsel review
    
    4. 💼 NEGOTIATION PRIORITIES:
       - Rank clauses from most critical to negotiate to acceptable as-is
       - Suggest bargaining positions and fallback options
    
    Present your analysis in a structured format that would help decision-makers understand the legal landscape of this opportunity.
    """
    
    # Use the retrieval QA chain to get relevant RFP sections and analyze them
    response = risk_qa_chain.invoke({"query": prompt})
    return response["result"]

def extract_criteria(state):
    """Tool function to extract mandatory eligibility criteria."""
    company_data = state.company_data
    criteria_qa_chain = state.criteria_qa_chain
    
    prompt = f"""
    You are an RFP Compliance Analyst. Your task is to extract and organize all the mandatory criteria and requirements from the RFP document.
    
    Company Profile:
    {json.dumps(company_data, indent=2)}
    
    Based on the RFP document, please:
    
    1. 📋 MANDATORY REQUIREMENTS:
       - Extract all "must-have" and "shall" requirements
       - Categorize by type (technical, business, legal, financial)
       - Note any disqualifying factors explicitly mentioned
    
    2. 🔢 EVALUATION CRITERIA:
       - List all scoring and evaluation criteria mentioned
       - Include point values or weightings if specified
       - Note minimum thresholds required to advance
    
    3. 📊 GAP ANALYSIS:
       - Compare company capabilities against requirements
       - Identify areas where the company fully meets criteria
       - Highlight gaps or areas needing strengthening
    
    4. 💡 COMPETITIVE POSITIONING:
       - Suggest areas where the company can demonstrate unique value
       - Identify requirements where the company exceeds expectations
       - Note potential differentiators from competitors
    
    Present this information in a structured format that helps the bid team understand exactly what is required for compliance and competitive advantage.
    """
    
    # Use the retrieval QA chain to get relevant RFP sections and analyze them
    response = criteria_qa_chain.invoke({"query": prompt})
    return response["result"]

class AgentState(BaseModel):
    query: str = None
    company_data: dict = None
    eligibility_qa_chain: RetrievalQA = None
    checklist_qa_chain: RetrievalQA = None
    risk_qa_chain: RetrievalQA = None
    criteria_qa_chain: RetrievalQA = None
    messages: List[BaseMessage] = Field(default_factory=list)
    tool_code_input: str = ""
    tool_call_id: str = ""
    tool_name: str = ""
    tool_input: Union[ToolInvocation, str] = ""
    tool_output: str = ""
    next_node: str = ""
    response: str = ""

    class Config:
        arbitrary_types_allowed = True

def user_input(state: AgentState):
    print("Entering user_input node")
    user_query = state.query
    return {"messages": [HumanMessage(content=user_query)]}

def agent_decision(state: AgentState):
    print("Entering agent_decision node")
    query = state.query.lower()
    if "eligibility" in query:
        print("Routing to eligibility_tool")
        return {"next_node": "eligibility_tool"}
    elif "checklist" in query:
        print("Routing to checklist_tool")
        return {"next_node": "checklist_tool"}
    elif "risk" in query:
        print("Routing to risk_tool")
        return {"next_node": "risk_tool"}
    elif "criteria" in query:
        print("Routing to criteria_tool")
        return {"next_node": "criteria_tool"}
    else:
        print("Routing to respond_to_user with default message")
        return {"next_node": "respond_to_user", "response": "No specific analysis requested. Please specify 'eligibility', 'checklist', 'risk', or 'criteria'."}

def eligibility_tool(state: AgentState):
    print("Entering eligibility_tool node")
    result = check_eligibility(state)
    return {"tool_output": str(result), "tool_name": "eligibility_tool"}

def checklist_tool(state: AgentState):
    print("Entering checklist_tool node")
    result = generate_checklist(state)
    return {"tool_output": str(result), "tool_name": "checklist_tool"}

def risk_tool(state: AgentState):
    print("Entering risk_tool node")
    result = analyze_risk(state)
    return {"tool_output": str(result), "tool_name": "risk_tool"}

def criteria_tool(state: AgentState):
    print("Entering criteria_tool node")
    result = extract_criteria(state)
    return {"tool_output": str(result), "tool_name": "criteria_tool"}

def format_tool_output(state: AgentState):
    print("Entering format_tool_output node")
    tool_output = state.tool_output
    tool_name = state.tool_name
    if tool_output and tool_name:
        print(f"Formatting tool output: {tool_output}")
        return {"messages": [FunctionMessage(content=tool_output, name=tool_name)]}
    return {}

def respond_to_user(state: AgentState):
    print("Entering respond_to_user node")
    if state.messages:
        last_message = state.messages[-1]
        print(f"Setting response from messages: {last_message.content}")
        return {"response": last_message.content}
    elif state.response:
        print(f"Setting response from state: {state.response}")
        return {"response": state.response}
    print("Setting default response")
    return {"response": "No analysis performed. Please try again with a specific query."}

def build_graph():
    builder = StateGraph(AgentState)

    builder.add_node("user_input", user_input)
    builder.add_node("agent_decision", agent_decision)
    builder.add_node("eligibility_tool", eligibility_tool)
    builder.add_node("checklist_tool", checklist_tool)
    builder.add_node("risk_tool", risk_tool)
    builder.add_node("criteria_tool", criteria_tool)
    builder.add_node("format_tool_output", format_tool_output)
    builder.add_node("respond_to_user", respond_to_user)

    builder.set_entry_point("user_input")
    builder.add_edge("user_input", "agent_decision")

    def route_decision(state: AgentState):
        next_node = state.next_node or "respond_to_user"
        print(f"Routing decision: {next_node}")
        return next_node

    builder.add_conditional_edges(
        "agent_decision",
        route_decision,
        {
            "eligibility_tool": "eligibility_tool",
            "checklist_tool": "checklist_tool",
            "risk_tool": "risk_tool",
            "criteria_tool": "criteria_tool",
            "respond_to_user": "respond_to_user"
        }
    )
    builder.add_edge("eligibility_tool", "format_tool_output")
    builder.add_edge("checklist_tool", "format_tool_output")
    builder.add_edge("risk_tool", "format_tool_output")
    builder.add_edge("criteria_tool", "format_tool_output")
    builder.add_edge("format_tool_output", "respond_to_user")
    builder.add_edge("respond_to_user", END)

    return builder.compile()

if __name__ == "__main__":
    company_data_file = r"data\Company Data.docx"
    rfp_file_path = r"data\ELIGIBLE RFP - 2.pdf"

    company_data = load_company_data(company_data_file)
    if not company_data:
        print("Exiting due to company data loading error.")
        exit()
    print("Loaded Company Data:", company_data)

    print(f"Parsing RFP document: {rfp_file_path}")
    rfp_text = parse_document_llama_parse(rfp_file_path)
    if not rfp_text:
        print("Exiting due to RFP parsing error.")
        exit()

    vector_store = setup_chroma_vector_store(hf_embeddings)
    if vector_store is None:
        print("Exiting due to Chroma setup error.")
        exit()

    vector_store = embed_and_store_in_chroma(vector_store, rfp_text)
    if vector_store is None:
        print("Exiting due to embedding error.")
        exit()

    eligibility_qa_chain = create_retrieval_qa_chain(vector_store)
    checklist_qa_chain = create_retrieval_qa_chain(vector_store)
    risk_qa_chain = create_retrieval_qa_chain(vector_store)
    criteria_qa_chain = create_retrieval_qa_chain(vector_store)

    app = build_graph()

    print("\nWelcome to the RFP Analysis Tool (LangGraph Version with Chroma)!")
    print("Using Groq LLM for text generation with Chroma for RAG.")
    print("Supports PDF RFP and DOCX Company Data documents.")
    print("You can ask questions related to:")
    print("- Eligibility: Check if your company meets all RFP requirements")
    print("- Checklist: Generate a comprehensive submission checklist")
    print("- Risk Analysis: Identify and assess contractual and legal risks")
    print("- Criteria Extraction: Extract and organize all mandatory criteria")
    print("Type 'exit' to quit.")

    while True:
        query = input("\nEnter your query (or type 'exit'): ")
        if query.lower() == 'exit':
            print("Exiting...")
            break

        inputs = {
            "query": query,
            "company_data": company_data,
            "eligibility_qa_chain": eligibility_qa_chain,
            "checklist_qa_chain": checklist_qa_chain,
            "risk_qa_chain": risk_qa_chain,
            "criteria_qa_chain": criteria_qa_chain,
            "messages": [],
        }
        output = app.invoke(inputs)
        if "response" in output:
            print(f"\nAI Response: {output['response']}")
        else:
            print("\nNo response generated. Final state:",output)